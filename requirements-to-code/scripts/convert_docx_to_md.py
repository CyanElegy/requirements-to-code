#!/usr/bin/env python3
"""
Convert a .docx file to markdown.

Usage:
    python convert_docx_to_md.py <input.docx> [output.md]

Tries pandoc first (best quality), then python-docx, then macOS textutil.
If output path is omitted, writes to <input-stem>.md.
"""

import sys
import subprocess
import os
from pathlib import Path


def try_pandoc(input_path: str, output_path: str) -> bool:
    """Try converting with pandoc. Returns True on success."""
    try:
        subprocess.run(
            ["pandoc", input_path, "-t", "markdown", "-o", output_path],
            check=True, capture_output=True, text=True,
        )
        return True
    except (FileNotFoundError, subprocess.CalledProcessError):
        return False


def try_python_docx(input_path: str, output_path: str) -> bool:
    """Try converting with python-docx. Returns True on success."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return False

    try:
        doc = Document(input_path)
        lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            # Detect heading style
            if para.style.name.startswith("Heading"):
                level = int(para.style.name.split()[-1])
                lines.append(f"{'#' * level} {text}")
            elif para.style.name.startswith("List"):
                lines.append(f"- {text}")
            else:
                lines.append(text)

        # Extract tables
        for table in doc.tables:
            lines.append("")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")

        Path(output_path).write_text("\n".join(lines), encoding="utf-8")
        return True
    except Exception:
        return False


def try_textutil(input_path: str, output_path: str) -> bool:
    """Try macOS textutil. Returns True on success."""
    try:
        subprocess.run(
            ["textutil", "-convert", "txt", input_path, "-output", output_path],
            check=True, capture_output=True,
        )
        # textutil outputs .txt; read and save as .md
        txt_path = str(Path(input_path).with_suffix(".txt"))
        if os.path.exists(txt_path):
            content = Path(txt_path).read_text(encoding="utf-8", errors="replace")
            Path(output_path).write_text(content, encoding="utf-8")
            os.remove(txt_path)
        return True
    except (FileNotFoundError, subprocess.CalledProcessError):
        return False


def main():
    if len(sys.argv) < 2:
        print("Usage: python convert_docx_to_md.py <input.docx> [output.md]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else str(Path(input_path).with_suffix(".md"))

    if not os.path.exists(input_path):
        print(f"Error: file not found: {input_path}")
        sys.exit(1)

    if not input_path.endswith(".docx"):
        print("Error: input must be a .docx file")
        sys.exit(1)

    for method_name, method_fn in [
        ("pandoc", try_pandoc),
        ("python-docx", try_python_docx),
        ("textutil", try_textutil),
    ]:
        if method_fn(input_path, output_path):
            print(f"Converted {input_path} → {output_path} (via {method_name})")
            sys.exit(0)

    print("Error: no conversion method available. Install pandoc, python-docx, or use macOS textutil.")
    sys.exit(1)


if __name__ == "__main__":
    main()
